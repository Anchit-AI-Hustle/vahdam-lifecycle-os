#!/usr/bin/env python3
"""Convert the VAHDAM PRD markdown into a brand-styled Word .docx.
Handles: #/##/### headings, **bold**, `code`, [links](url), bullet + numbered
lists, GitHub tables, --- rules, blockquotes, paragraphs. Brand palette only."""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOREST = RGBColor(0x00, 0x4A, 0x2B)
GOLD   = RGBColor(0xAB, 0x87, 0x43)
INK    = RGBColor(0x17, 0x17, 0x17)
CREAM_HEX = "FBF5EA"
FOREST_HEX = "004A2B"

SRC, OUT = sys.argv[1], sys.argv[2]
doc = Document()
# Base style
normal = doc.styles["Normal"]
normal.font.name = "Helvetica"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hex_fill)
    tcPr.append(sh)

INLINE = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))')
def add_runs(paragraph, text):
    # strip stray markdown emphasis underscores kept; handle bold/code/links
    parts = INLINE.split(text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            r = paragraph.add_run(p[2:-2]); r.bold = True
        elif p.startswith("`") and p.endswith("`"):
            r = paragraph.add_run(p[1:-1]); r.font.name = "Consolas"; r.font.color.rgb = FOREST
        elif p.startswith("*") and p.endswith("*") and len(p) > 2:
            r = paragraph.add_run(p[1:-1]); r.italic = True
        elif p.startswith("[") and "](" in p:
            label = p[1:p.index("](")]
            r = paragraph.add_run(label); r.font.color.rgb = GOLD; r.underline = True
        else:
            paragraph.add_run(p)

lines = open(SRC, encoding="utf-8").read().split("\n")
i = 0
n = len(lines)
first_h1 = True
while i < n:
    line = lines[i]
    # Table block
    if line.strip().startswith("|") and i + 1 < n and re.match(r'^\s*\|[\s:\-|]+\|\s*$', lines[i+1]):
        header = [c.strip() for c in line.strip().strip("|").split("|")]
        rows = []
        i += 2
        while i < n and lines[i].strip().startswith("|"):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
            i += 1
        table = doc.add_table(rows=1, cols=len(header))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for j, h in enumerate(header):
            hdr[j].text = ""
            para = hdr[j].paragraphs[0]
            add_runs(para, h)
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(0xFB,0xF5,0xEA); run.font.size = Pt(9.5)
            shade(hdr[j], FOREST_HEX)
        for r in rows:
            cells = table.add_row().cells
            for j in range(len(header)):
                cells[j].text = ""
                para = cells[j].paragraphs[0]
                add_runs(para, r[j] if j < len(r) else "")
                for run in para.runs:
                    run.font.size = Pt(9.5)
        doc.add_paragraph()
        continue
    # Headings
    m = re.match(r'^(#{1,6})\s+(.*)$', line)
    if m:
        level = len(m.group(1)); text = m.group(2)
        if level == 1:
            h = doc.add_paragraph(); h.space_after = Pt(6)
            run = h.add_run(re.sub(r'\*\*|`', '', text)); run.bold = True
            run.font.size = Pt(22 if first_h1 else 16); run.font.color.rgb = FOREST
            first_h1 = False
        else:
            h = doc.add_paragraph()
            run = h.add_run(re.sub(r'\*\*|`', '', text)); run.bold = True
            run.font.size = Pt({2:15, 3:12}.get(level, 11)); run.font.color.rgb = FOREST if level==2 else GOLD
        i += 1; continue
    # Horizontal rule
    if line.strip() == "---":
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "AB8743")
        pbdr.append(bottom); pPr.append(pbdr)
        i += 1; continue
    # Blockquote
    if line.strip().startswith(">"):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
        add_runs(p, line.strip()[1:].strip())
        for run in p.runs: run.italic = True; run.font.color.rgb = FOREST
        i += 1; continue
    # Lists
    mlist = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)$', line)
    if mlist:
        indent = len(mlist.group(1)); ordered = bool(re.match(r'\d+\.', mlist.group(2)))
        style = "List Number" if ordered else "List Bullet"
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph(); p.add_run("• ")
        p.paragraph_format.left_indent = Inches(0.25 + 0.25*(indent//2))
        add_runs(p, mlist.group(3))
        i += 1; continue
    # Blank
    if not line.strip():
        i += 1; continue
    # Paragraph
    p = doc.add_paragraph()
    add_runs(p, line)
    i += 1

doc.save(OUT)
print("wrote", OUT)
